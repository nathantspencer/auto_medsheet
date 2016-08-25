from bs4 import BeautifulSoup
import code
from docx import Document
from docx.shared import Pt
import requests
import re
import sys

# pdr_surf function
def pdr_surf(file_name, medications):

    # these little lists will hold our grabbed data
    med = []
    gen = []
    cla = []
    moa = []
    nsg = []
    adv = []

    # for each medication passed...
    for medication in medications:

        # search all pages of drugs starting with that letter
        search_url = 'http://www.pdr.net/search-results?q=' + \
            medication.lower()
        response = requests.get(search_url)
        soup = BeautifulSoup(response.text, 'lxml')
        pea_soup = soup.findAll('a', href=re.compile(medication.lower() + \
            '\?druglabelid'), text=re.compile(medication.title()))

        # if found, grab the link to the drug page
        if len(pea_soup) > 0:
            print('\nFound ' + medication.title() + ' on PDR.net!')
            soup_string = str(pea_soup.pop())
            href = re.search('a href="(.*?)"', soup_string)
            med_link = href.group(1)

        # otherwise, break out of the loop
        else:
            print('\nWARNING: Could not find ' + medication.title() + \
                ' on PDR.net!')

            # since we couldn't find the drug, we'll leave its data blank
            med.append('')
            gen.append('')
            cla.append('')
            moa.append('')
            nsg.append('')
            adv.append('')
            continue;

        # otherwise, navigate to the drug summary page
        response = requests.get(med_link);
        soup = BeautifulSoup(response.text, 'lxml')
        pea_soup = soup.findAll('a', text=re.compile('Drug Summary'))
        soup_string = str(pea_soup.pop())
        href = re.search('a href="(.*?)"', soup_string)
        drug_summary_link = href.group(1)
        response = requests.get(drug_summary_link)

        # parse drug summary page for...
        soup = BeautifulSoup(response.text, 'lxml')

        # medication name
        medication_name = medication.title()
        med.append(medication_name)

        # generic name
        generic_name = ''
        pea_soup = soup.findAll('div', { "class" : "drugSummaryLabel" }, \
            text=re.compile('\((.*?)\)'))
        if(len(pea_soup) > 0):
            print('Found generic name for ' + medication.title() + '.')
            soup_string = str(pea_soup.pop())
            generic_name = re.search('\((.*?)\)', soup_string)
            generic_name = generic_name.group(1)
        else:
            print('WARNING: Could not find generic name for ' + \
                medication.title() + '.')
        gen.append(generic_name)

        # therapeutic class
        therapeutic_class = ''
        pea_soup = soup.findAll('h3', { "class" : "drugSummary" }, \
            text=re.compile('THERAPEUTIC CLASS'))
        if(len(pea_soup) > 0):
            print('Found therapeutic class for ' + medication.title() + '.')
            therapeutic_class = pea_soup.pop().nextSibling.nextSibling.text
        else:
            print('WARNING: Could not find therapeutic class for ' + \
                medication.title() + '.')
        cla.append(therapeutic_class)

        # mechanism of action
        mech_of_action = ''
        pea_soup = soup.findAll('h3', { "class" : "drugSummary" }, \
            text=re.compile('MECHANISM OF ACTION'))
        if(len(pea_soup) > 0):
            print('Found mechanism of action for ' + medication.title() + '.')
            mech_of_action = pea_soup.pop().nextSibling.nextSibling.text
        else:
            print('WARNING: Could not find mechanism of action for ' + \
                medication.title() + '.')
        moa.append(mech_of_action)

        # nursing assessment
        assessment = ''
        pea_soup = soup.findAll('h3', { "class" : "drugSummary" }, \
            text=re.compile('ASSESSMENT'))
        if(len(pea_soup) > 0):
            print('Found nursing assessment for ' + medication.title() + '.')
            assessment = pea_soup.pop().nextSibling.nextSibling.text
        else:
            print('WARNING: Could not find nursing assessment for ' + \
                medication.title() + '.')
        nsg.append(assessment)

        # adverse reactions
        adverse_reactions = ''
        pea_soup = soup.findAll('h3', { "class" : "drugSummary" }, \
            text=re.compile('ADVERSE REACTIONS'))
        if(len(pea_soup) > 0):
            print('Found adverse reactions for ' + medication.title() + '.')
            adverse_reactions = pea_soup.pop().nextSibling.nextSibling.text
        else:
            print('WARNING: Could not find adverse reactions for ' + \
                medication.title() + '.')
        adv.append(adverse_reactions)

    # all done, now lets write to a document
    print('\nWriting information to ' + file_name + '.docx ...')
    d = Document()

    # we'll start off with our headers
    table = d.add_table(rows=len(medications)+1, cols=8, style=d.styles['Table Grid'])
    table.style.font.size = Pt(7)
    table.cell(0, 0).text = 'Med.'
    table.cell(0, 1).text = 'Generic'
    table.cell(0, 2).text = 'Dose'
    table.cell(0, 3).text = 'Class'
    table.cell(0, 4).text = 'MOA'
    table.cell(0, 5).text = 'Nsg.'
    table.cell(0, 6).text = 'Why'
    table.cell(0, 7).text = 'Other'

    # make those headers bold
    for i in range(8):
        table.cell(0,i).paragraphs[0].runs[0].bold = True

    # fill in the text fields
    for i in range(1, len(medications)+1):
        table.cell(i, 0).text = med[i-1]
        table.cell(i, 1).text = gen[i-1]
        table.cell(i, 3).text = cla[i-1]
        table.cell(i, 4).text = moa[i-1]
        table.cell(i, 5).text = nsg[i-1]
        table.cell(i, 7).text = 'Adverse Effects: ' + adv[i-1]

    # save the file and we're all done!
    d.save(file_name + '.docx')
    print('PDR information written to ' + file_name + '.docx successfully!')
    print('Good luck in your clinical!\n\t -- Nathan\n')


# help text and launch of pdr_surf
if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('PDR_SURFER -- Written by Nathan Spencer 2016')
        print('Usage: python pdr_surfer.py [fileNameWithoutExtension] [drugName] [anotherDrug]')
    else:
        pdr_surf(sys.argv[1], sys.argv[2:])
